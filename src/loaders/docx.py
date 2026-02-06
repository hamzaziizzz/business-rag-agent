from __future__ import annotations

"""DOCX loader for document ingestion."""

from io import BytesIO

from src.rag.types import Document


class DocxLoaderError(RuntimeError):
    """Raised when DOCX loading fails."""
    pass


def load_docx_bytes(data: bytes, doc_id: str, source: str) -> Document:
    """Load a DOCX file from bytes into a Document."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocxLoaderError("python-docx is required to load DOCX files") from exc

    doc = DocxDocument(BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n".join(parts).strip()
    return Document(doc_id=doc_id, content=content, metadata={"source": source})
