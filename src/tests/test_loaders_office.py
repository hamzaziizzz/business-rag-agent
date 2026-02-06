from __future__ import annotations

"""Loader tests for office document formats."""

from io import BytesIO

from src.loaders.csv_loader import load_csv_bytes
from src.loaders.docx import load_docx_bytes
from src.loaders.xlsx import load_xlsx_bytes


def test_load_csv_bytes() -> None:
    """Ensure CSV bytes load into Document content."""
    data = b"col1,col2\n1,2\n"
    doc = load_csv_bytes(data, doc_id="csv-1", source="data.csv")
    assert "col1" in doc.content
    assert "1,2" in doc.content


def test_load_docx_bytes() -> None:
    """Ensure DOCX bytes load into Document content."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return
    docx = DocxDocument()
    docx.add_paragraph("Quarterly report")
    docx.add_paragraph("Revenue increased.")
    buffer = BytesIO()
    docx.save(buffer)
    doc = load_docx_bytes(buffer.getvalue(), doc_id="docx-1", source="report.docx")
    assert "Quarterly report" in doc.content
    assert "Revenue increased." in doc.content


def test_load_xlsx_bytes() -> None:
    """Ensure XLSX bytes load into Document content."""
    try:
        from openpyxl import Workbook
    except ImportError:
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "value"])
    ws.append(["Ada", 10])
    buffer = BytesIO()
    wb.save(buffer)
    doc = load_xlsx_bytes(buffer.getvalue(), doc_id="xlsx-1", source="data.xlsx")
    assert "Sheet: Sheet1" in doc.content
    assert "Ada" in doc.content
